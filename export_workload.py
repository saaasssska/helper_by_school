import sys
# -*- coding: utf-8 -*-
import sqlite3
from docx import Document
from PyQt5 import uic
from PyQt5.QtWidgets import QApplication, QMainWindow, QMessageBox, QDialog, QFileDialog


class MyExport_workload(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('frontend\export_workload.ui', self)
        self.con = sqlite3.connect('db_subjects.db')
        self.push_export.clicked.connect(self.export)

    def export(self):
        filename = self.line_filename.text()
        if filename.isspace():
            pass
        cur = self.con.cursor()
        teachers = [" ".join(i) for i in cur.execute("""SELECT 
                                            teachers.surname,
                                            teachers.name,
                                            teachers.patronymic
                                            FROM
                                            teachers
                                            INNER JOIN teacher_workload
                                            ON teachers.id_teacher = teacher_workload.id_teacher;""")]
        document = Document()
        for i in teachers:
            full_name = i.split()
            surname, name, patronymic = full_name[0], full_name[1], full_name[2]
            id_teacher = [i[0] for i in cur.execute("""SELECT id_teacher from teachers
                                            WHERE surname = ? AND name = ? AND patronymic = ?""",
                                                    [surname, name, patronymic])]
            id_teacher = id_teacher[0]
            workload = [list(i) for i in cur.execute("""SELECT class, build, subject, hours FROM teacher_workload
                                            WHERE id_teacher = ?""", [id_teacher])]
            all_hours = 0
            for j in workload:
                all_hours += int(j[-1])
            document.add_heading(i, 0)
            for elem in workload:
                document.add_paragraph("Класс: " + elem[0] + ', школа: ' + elem[1]
                                           + ', предмет: ' + elem[2] + ', часы: ' + elem[3])
            document.add_paragraph("Общее количество часов: " + str(all_hours))
            document.add_page_break()
        fname = QFileDialog.getExistingDirectory()
        filename = str(filename) + '.docx'
        if fname.isspace():
            pass
        file = fname + '/' + filename
        document.save(file)
        self.label_save.setText("Файлы успешно экспортированы")
        self.line_filename.clear()



